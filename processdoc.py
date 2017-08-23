#This module is for processing MS Word documents (For docx written in MS Word 2007 and above)

from textblob import TextBlob
from docx import Document

class ProcessDoc:
    def __init__(self, file_name):
        self.document = Document(file_name)


    def no_of_paragraphs(self):
        '''Returns number of paragraphs in the document'''
        return len(self.document.paragraphs)


    def specific_paragraph(self, number:int):
        '''Returns a paragraph of choice given its number/position in the document'''
        needed_paragraph = self.document.paragraphs[number - 1]
        return needed_paragraph.text

    def all_paragraphs(self):
        '''Returns a list of all the paragraphs in a document, with each paragraph converted to a textblob'''
        paragraphs = list()
        for para in self.document.paragraphs:
            paragraphs.append(TextBlob(para.text))
        return paragraphs

    def convert_to_text(self):
        '''Converts the text from the document into a Text'''
        text = ''
        for para in self.document.paragraphs:
            text += para.text
        return text

    def convert_to_textblob(self):
        '''Converts the text from the document into a TextBlob'''
        text = ''
        for para in self.document.paragraphs:
            text += para.text
        return TextBlob(text)

class Textblob_with_file_name(TextBlob):
    def __init__(self, file_name,tokenizer=None,pos_tagger=None,np_extractor=None,analyzer=None,parser=None,classifier=None,clean_html:bool = False):
        self.document = ProcessDoc(file_name)
        TextBlob.__init__(self,str(self.document.convert_to_text()),tokenizer,pos_tagger,np_extractor,analyzer,parser,classifier,clean_html)

# if __name__ == '__main__':
#     ##For testing
#     document = Textblob_with_file_name('AP Final Draft - Meghana Jain.docx').document
#     print(document.no_of_paragraphs())
#     print(document.all_paragraphs())
#     text = document.convert_to_textblob()
#     print(text.words)
#     print(text.sentences) # It is counting the title in the first sentence, Is that supposed to happen??? Thats just the way my document is formatted..i'll look into it
